#!/usr/bin/env python3
"""
Generate a comprehensive game analysis document for The King in Yellow.
Uses the game's Lovecraftian color palette for styling.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ═══════════════════════════════════════════
# GAME COLOR PALETTE
# ═══════════════════════════════════════════

VOID = RGBColor(26, 10, 46)
SHADOW = RGBColor(45, 27, 78)
ELDRITCH = RGBColor(74, 14, 78)
YELLOW = RGBColor(212, 160, 23)
AMBER = RGBColor(184, 134, 11)
BONE = RGBColor(212, 197, 169)
CRIMSON = RGBColor(139, 0, 0)
BLOOD = RGBColor(196, 30, 58)
MIST = RGBColor(143, 188, 143)
FROST = RGBColor(70, 130, 180)
EMBER = RGBColor(255, 69, 0)
MADNESS_COLOR = RGBColor(255, 99, 71)
GOLD = RGBColor(241, 196, 15)
ASH = RGBColor(105, 105, 105)
DARK_BG = RGBColor(12, 6, 24)
PANEL_BG = RGBColor(20, 12, 38)
HP_GREEN = RGBColor(46, 204, 113)
HP_RED = RGBColor(231, 76, 60)
XP_PURPLE = RGBColor(155, 89, 182)
GOLD_TRIM = RGBColor(218, 165, 32)
ELDRITCH_GOLD = RGBColor(232, 185, 45)
ELDRITCH_PURPLE = RGBColor(175, 130, 225)
OBSIDIAN = RGBColor(16, 12, 28)
OBSIDIAN_MID = RGBColor(28, 20, 45)
WHITE = RGBColor(255, 255, 255)
OFF_WHITE = RGBColor(240, 235, 225)

# Class colors
SCHOLAR_COLOR = RGBColor(180, 100, 230)
BRUTE_COLOR = RGBColor(220, 50, 50)
WARDEN_COLOR = FROST
SHADOWBLADE_COLOR = MIST
PROPHET_COLOR = MADNESS_COLOR


def set_cell_bg(cell, color: RGBColor):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.hex()}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_bg_hex(cell, hex_color: str):
    """Set cell background color from hex string."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, font_name="Calibri", font_size=11,
                         color=OFF_WHITE, bold=False, italic=False,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                         space_before=0):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_mixed_paragraph(doc, segments, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_after=6, space_before=0):
    """Add a paragraph with mixed formatting.
    segments: list of (text, font_name, font_size, color, bold, italic)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for seg in segments:
        text, font_name, font_size, color, bold, italic = seg
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic
    return p


def add_divider(doc, color=GOLD_TRIM):
    """Add a horizontal gold divider line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = color
    return p


def add_section_header(doc, text, level=1):
    """Add a section header with Lovecraftian styling."""
    if level == 1:
        p = add_styled_paragraph(doc, text.upper(), font_size=16,
                                 color=ELDRITCH_GOLD, bold=True,
                                 space_before=18, space_after=8)
        add_divider(doc, GOLD_TRIM)
    elif level == 2:
        p = add_styled_paragraph(doc, text, font_size=13,
                                 color=YELLOW, bold=True,
                                 space_before=14, space_after=6)
    elif level == 3:
        p = add_styled_paragraph(doc, text, font_size=11,
                                 color=ELDRITCH_PURPLE, bold=True,
                                 space_before=10, space_after=4)
    return p


def add_bullet(doc, text, indent=0, color=OFF_WHITE, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(1.0 + indent * 0.8)

    if bold_prefix:
        run_bullet = p.add_run("▸ ")
        run_bullet.font.size = Pt(10)
        run_bullet.font.color.rgb = GOLD_TRIM
        run_bold = p.add_run(bold_prefix)
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = YELLOW
        run_bold.font.bold = True
        run_text = p.add_run(text)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = color
    else:
        run = p.add_run(f"▸ {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = color
    return p


def add_impact_difficulty_table(doc, items):
    """Add a table with Impact and Difficulty columns.
    items: list of (name, impact, difficulty, description)
    Impact/Difficulty: ★ to ★★★★★
    """
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header
    header_cells = table.rows[0].cells
    headers = ["Improvement", "Impact", "Difficulty", "Notes"]
    for i, (cell, header) in enumerate(zip(header_cells, headers)):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for item_name, impact, difficulty, notes in items:
        row_cells = table.add_row().cells
        # Name
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(item_name)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE
        run.font.bold = True

        # Impact
        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(impact)
        run.font.size = Pt(9)
        run.font.color.rgb = HP_GREEN if "★★★★" in impact else YELLOW if "★★★" in impact else AMBER

        # Difficulty
        set_cell_bg_hex(row_cells[2], "140C26")
        p = row_cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(difficulty)
        run.font.size = Pt(9)
        run.font.color.rgb = HP_RED if "★★★★" in difficulty else YELLOW if "★★★" in difficulty else MIST

        # Notes
        set_cell_bg_hex(row_cells[3], "140C26")
        p = row_cells[3].paragraphs[0]
        run = p.add_run(notes)
        run.font.size = Pt(8)
        run.font.color.rgb = ASH

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(2.5)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(6.5)

    return table


def create_analysis_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = OFF_WHITE

    # Set page background color (dark)
    bg_elem = parse_xml(
        f'<w:background {nsdecls("w")} w:color="0C0618"/>'
    )
    doc.element.insert(0, bg_elem)

    # Also set the display background
    settings = doc.settings.element
    display_bg = parse_xml(
        f'<w:displayBackgroundShape {nsdecls("w")}/>'
    )
    settings.append(display_bg)

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════

    # Spacer
    for _ in range(4):
        add_styled_paragraph(doc, "", font_size=12)

    add_styled_paragraph(doc, "⛧", font_size=36, color=GOLD_TRIM,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "THE KING IN YELLOW", font_size=28,
                         color=ELDRITCH_GOLD, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "A Lovecraftian Dungeon Crawler", font_size=14,
                         color=ELDRITCH_PURPLE, italic=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_divider(doc, GOLD_TRIM)
    add_styled_paragraph(doc, "COMPREHENSIVE ANALYSIS", font_size=18,
                         color=YELLOW, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
    add_styled_paragraph(doc, "GAME STATUS · SECTOR BREAKDOWN · IMPROVEMENT ROADMAP",
                         font_size=10, color=AMBER,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_divider(doc, GOLD_TRIM)

    for _ in range(3):
        add_styled_paragraph(doc, "", font_size=12)

    add_styled_paragraph(doc, "Generated: 2026-04-23", font_size=9,
                         color=ASH, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Engine: Python + Pygame  |  1280×720  |  5 Classes  |  20 Floors",
                         font_size=9, color=ASH,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════

    add_section_header(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Game Overview & Architecture",
        "3. Sector Analysis",
        "   3.1  Architecture & Code Quality",
        "   3.2  Combat System",
        "   3.3  Class Design & Balance",
        "   3.4  Visual & UI Design",
        "   3.5  Content & Narrative",
        "   3.6  Audio & Polish",
        "   3.7  Save System & Persistence",
        "   3.8  Testing & CI/CD",
        "4. Master Improvement List",
        "5. Priority Roadmap",
        "6. Appendix: Technical Metrics",
    ]
    for item in toc_items:
        add_styled_paragraph(doc, item, font_size=10, color=BONE,
                             space_after=3)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════

    add_section_header(doc, "1. Executive Summary", level=1)

    add_styled_paragraph(doc,
        "The King in Yellow is a turn-based Lovecraftian dungeon crawler built with Python and Pygame. "
        "Inspired by Buried Bornes, it features 5 unique classes, a deep buff/debuff system with 67+ status "
        "effects, procedural item generation, and a richly styled gothic UI with parchment textures, "
        "ornate gold frames, and particle effects.",
        color=OFF_WHITE, space_after=8)

    add_styled_paragraph(doc,
        "The project has reached a strong foundation: 1,156 automated tests, a CI/CD pipeline, "
        "dependency injection architecture, surface pooling for performance, and a well-organized "
        "codebase split across 50+ Python files. The game is fully playable through 20 floors with "
        "a boss encounter, 5 distinct classes each with ~40 skills, and a save/load system.",
        color=OFF_WHITE, space_after=8)

    add_section_header(doc, "Overall Assessment", level=2)

    # Grade box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("CODE QUALITY GRADE:  A−  (88/100)")
    run.font.size = Pt(14)
    run.font.color.rgb = ELDRITCH_GOLD
    run.font.bold = True

    add_styled_paragraph(doc,
        "Strengths: Excellent test coverage, clean modular architecture, comprehensive combat "
        "system, data-driven design, and strong visual identity. The codebase is well-positioned "
        "for the planned combat overhaul (limb loss, madness rework, curse system, religion system).",
        color=OFF_WHITE, space_after=6)

    add_styled_paragraph(doc,
        "Key Gaps: No audio system, limited enemy variety (6 + boss), thin content (6 events, 3 traps), "
        "no meta-progression or run-to-run persistence, and several planned-but-unimplemented systems "
        "(limbs, curses, religion, skill evolution).",
        color=OFF_WHITE, space_after=6)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 2. GAME OVERVIEW & ARCHITECTURE
    # ═══════════════════════════════════════════

    add_section_header(doc, "2. Game Overview & Architecture", level=1)

    add_section_header(doc, "File Structure", level=2)

    arch_items = [
        ("pygame_game.py", "Main game class, loop, screen transitions (~237 lines)"),
        ("save_system.py", "JSON save/load with 5 slots, version 2 format"),
        ("shared/", "Constants, assets, rendering, surface pool, game context (4 modules)"),
        ("engine/", "Combat, damage, skills, status effects, items, world (7 modules)"),
        ("data/", "Classes (5), enemies, items, events, narratives, constants, JSON data"),
        ("screens/", "17 screen modules (combat split into package with particles & renderer)"),
        ("config/", "Settings manager with dot-notation access, JSON config file"),
        ("tests/", "1,156 tests across 5 test suites + unified runner"),
    ]
    for name, desc in arch_items:
        add_mixed_paragraph(doc, [
            ("▸ ", "Calibri", 10, GOLD_TRIM, False, False),
            (name, "Consolas", 9, YELLOW, True, False),
            (f" — {desc}", "Calibri", 10, OFF_WHITE, False, False),
        ], space_after=3)

    add_section_header(doc, "Key Metrics", level=2)

    metrics = [
        ("Total Python Files", "50+"),
        ("Lines of Code", "~8,700"),
        ("Automated Tests", "1,156"),
        ("Playable Classes", "5 (Scholar, Brute, Warden, Shadowblade, Mad Prophet)"),
        ("Skills per Class", "~40 (201 total)"),
        ("Buff/Debuff Types", "67+"),
        ("Enemy Types", "6 + 1 Boss (Hastur)"),
        ("Game Floors", "20 (with boss on floor 20)"),
        ("Equipment Slots", "6 (weapon, armor, accessory, boots, 2 rings)"),
        ("Save Slots", "5"),
        ("CI Pipeline", "GitHub Actions (lint + test + coverage)"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, ["Metric", "Value"]):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for metric, value in metrics:
        row_cells = table.add_row().cells
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(metric)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE
        run.font.bold = True

        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = OFF_WHITE

    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 3. SECTOR ANALYSIS
    # ═══════════════════════════════════════════

    add_section_header(doc, "3. Sector Analysis", level=1)

    # ── 3.1 Architecture ──
    add_section_header(doc, "3.1  Architecture & Code Quality", level=2)

    add_styled_paragraph(doc, "Status: ✅ Strong", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The codebase follows a clean layered architecture: data → engine → shared → screens → main. "
        "Dependency injection via GameContext decouples screens from the Game class, enabling unit testing "
        "without Pygame. The codebase has been through multiple refactoring passes, splitting monolithic "
        "files into well-organized packages.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Completed Refactors", level=3)
    completed = [
        "game_data.py (650 lines) → data/ package (6 files)",
        "game_engine.py (1,361 lines) → engine/ package (7 modules)",
        "pygame_game.py (3,225 lines) → shared/ + screens/ + main (237 lines)",
        "engine/combat.py (1,275 lines) → combat.py + damage.py + skills.py + items.py",
        "Dependency Injection: 145 self.game.* → self.ctx.* across all 17 screens",
        "Surface pooling: zero per-frame Surface allocations in draw paths",
        "Parchment texture full caching: zero per-frame allocation on cache hit",
    ]
    for item in completed:
        add_bullet(doc, item, color=BONE)

    add_section_header(doc, "Remaining Issues", level=3)
    issues = [
        ("Naming conventions: ", "31 private helpers use camelCase instead of snake_case in engine/skills.py"),
        ("Division-by-zero: ", "4 instances in damage.py and skills.py (low risk but should be guarded)"),
        ("Docstring coverage: ", "26.3% of functions have docstrings (target: 80%+)"),
        ("Type safety: ", "17 mypy warnings (Dict indexing, non-critical)"),
        ("Magic numbers: ", "Some hardcoded values in engine/items.py and engine/world.py"),
    ]
    for prefix, text in issues:
        add_bullet(doc, text, bold_prefix=prefix, color=OFF_WHITE)

    add_impact_difficulty_table(doc, [
        ("Guard division-by-zero", "★★★★", "★", "30 min — prevents rare crashes"),
        ("Rename camelCase helpers", "★★", "★★", "1 hour — improves consistency"),
        ("Add docstrings to engine", "★★★", "★★★", "3-4 hours — critical for maintainability"),
        ("Add TypedDict for data structures", "★★★", "★★", "2 hours — fixes 17 mypy warnings"),
        ("Extract remaining magic numbers", "★★", "★★", "1 hour — aids balancing"),
    ])

    doc.add_page_break()

    # ── 3.2 Combat System ──
    add_section_header(doc, "3.2  Combat System", level=2)

    add_styled_paragraph(doc, "Status: ✅ Very Strong — Deep and well-tested", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The turn-based combat system is the game's crown jewel. It features a sophisticated damage "
        "formula with physical/magic/curse types, defense percentage reduction (df/(df+50)), crit system, "
        "accuracy/miss chance, shields, barriers, evasion, lifesteal, armor pierce, and a 67+ buff/debuff "
        "registry. Enemy intent is pre-shown each turn, enabling strategic play. Boss fights have 3 phases "
        "with escalating abilities. Victory animations feature a 4-phase sequence.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Strengths", level=3)
    strengths = [
        "Data-driven buff/debuff system — adding new effects = one function + one dict entry",
        "Handler registry pattern for heals (15 types), shields (10 types), buffs (28 types)",
        "Damage preview on skill hover shows ~dmg range after enemy defense",
        "271 combat tests covering all 5 classes, damage, buffs, debuffs, AI, boss phases",
        "Property-based testing with Hypothesis (515 tests) for combat balance invariants",
        "Enemy AI with intent telegraphing — players can plan defensively",
    ]
    for s in strengths:
        add_bullet(doc, s, color=BONE)

    add_section_header(doc, "Gaps", level=3)
    gaps = [
        ("No limb system: ", "Planned but not implemented — would add strategic depth"),
        ("Madness is shallow: ", "Currently just a death timer at 100 — needs threshold effects"),
        ("No curse system: ", "Planned limb/body/essence curses would add counterplay"),
        ("Only 6 enemy types: ", "Gets repetitive by mid-game — needs floor-specific enemies"),
        ("Boss is only Hastur: ", "Need 2-3 more bosses for variety across runs"),
        ("Flee is binary: ", "Success/fail with no nuance — could add chase mechanics"),
    ]
    for prefix, text in gaps:
        add_bullet(doc, text, bold_prefix=prefix, color=OFF_WHITE)

    add_impact_difficulty_table(doc, [
        ("Limb Loss System", "★★★★★", "★★★★", "Major feature — new engine/limbs.py + GameState changes"),
        ("Madness Overhaul", "★★★★★", "★★★", "Threshold system with escalating effects per tier"),
        ("Curse System", "★★★★", "★★★", "New engine/curses.py — limb/body/essence types"),
        ("More Enemy Types (6→12+)", "★★★★", "★★", "Data-driven — just JSON + sprite assets"),
        ("Second Boss (e.g., Azathoth)", "★★★★", "★★★", "New boss data + unique phase mechanics"),
        ("Enemy Scaling per Floor", "★★★", "★", "Already have floor scaling — just needs tuning"),
        ("Flee Rework", "★★", "★★", "Add chase sequence, partial escape options"),
    ])

    doc.add_page_break()

    # ── 3.3 Class Design ──
    add_section_header(doc, "3.3  Class Design & Balance", level=2)

    add_styled_paragraph(doc, "Status: ✅ Strong — 5 distinct identities", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "Each class has a clear identity, primary stat, and ~40 skills spread across 3 tiers. "
        "The Scholar is a glass-cannon mage, Brute a berserker tank, Warden a paladin/healer, "
        "Shadowblade a rogue/assassin, and Mad Prophet a chaos gambler. Debuff distribution is "
        "well-differentiated — Mad Prophet has the most debuffs (12), while Brute focuses on "
        "physical power with fewer debuffs.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Class Balance Assessment", level=3)

    classes = [
        ("Scholar", SCHOLAR_COLOR, "INT", "70+6/lv", "41 skills",
         "Glass cannon archetype works well. Highest INT base. Dual-stat scaling (INT+WIS) on "
         "late skills adds depth. 12 debuff skills give strong crowd control. Doom execute is "
         "powerful but balanced by 3-turn delay."),
        ("Brute", BRUTE_COLOR, "STR", "110+10/lv", "40 skills",
         "Berserker fantasy is clear. Missing-HP scaling (Last Stand, Blood Frenzy) creates exciting "
         "low-HP gameplay. Lifesteal and Undying Fury keep you alive. Simple rotation but satisfying. "
         "Beginner-friendly."),
        ("Warden", WARDEN_COLOR, "WIS", "90+8/lv", "40 skills",
         "Shield-to-damage conversion (Retribution, Sacrifice Curse) is a unique mechanic. Best "
         "sustain in the game. Seal of Carcosa (T3 ultimate) is overloaded — full heal + massive "
         "shield + barrier + damage in one skill. May need tuning."),
        ("Shadowblade", SHADOWBLADE_COLOR, "AGI", "75+7/lv", "40 skills",
         "Poison stacking (up to 20% HP/turn) is the signature mechanic and feels great. Evasion "
         "stacking with Smoke Screen + Flicker creates a dance-on-death playstyle. Phantom Blades "
         "(4 hits, 100% armor pierce) is extremely powerful late-game."),
        ("Mad Prophet", PROPHET_COLOR, "WIS/LUCK", "80+7/lv", "40 skills",
         "Chaos fantasy is well-realized. Coin flips and gamble mechanics create exciting moments. "
         "Madness-immune skill (Madness Mastery) is a key differentiator. Most debuffs of any class. "
         "Hastur's Decree (triple debuff ultimate) is devastating."),
    ]

    for name, color, stat, hp, skills, analysis in classes:
        add_mixed_paragraph(doc, [
            ("▸ ", "Calibri", 10, GOLD_TRIM, False, False),
            (name, "Calibri", 11, color, True, False),
            (f"  |  Primary: {stat}  |  HP: {hp}  |  {skills}", "Calibri", 9, ASH, False, False),
        ], space_after=2)
        add_styled_paragraph(doc, analysis, font_size=9, color=OFF_WHITE, space_after=8)

    add_impact_difficulty_table(doc, [
        ("Skill Evolution / Mastery", "★★★★★", "★★★★", "Skills level up with use — adds meta-progression"),
        ("Skill Branching", "★★★★", "★★★", "At certain levels, choose between 2 upgrade paths"),
        ("Class-specific Madness", "★★★★", "★★★", "Each class reacts differently to madness"),
        ("Starting Loadout Choice", "★★★", "★", "Pick 2 of 4 starting skills — adds replayability"),
        ("Class Unlock System", "★★★", "★★", "Unlock classes through achievements, not all at once"),
    ])

    doc.add_page_break()

    # ── 3.4 Visual & UI ──
    add_section_header(doc, "3.4  Visual & UI Design", level=2)

    add_styled_paragraph(doc, "Status: ✅ Very Strong — Cohesive and atmospheric", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The visual identity is one of the game's greatest assets. The Lovecraftian palette of dark "
        "purples, golds, and bone whites creates an oppressive, ornate atmosphere. Procedural parchment "
        "textures with ornate gold frames, Cinzel Decorative Victorian fonts, cached glow text rendering, "
        "particle effects (dust, eldritch energy, blood), animated hover effects, screen shake, and a "
        "madness vignette overlay all contribute to a cohesive aesthetic.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "What's Working", level=3)
    working = [
        "Parchment texture system with ornate gold frames on all 17 screens",
        "Cinzel Decorative (titles) + Cinzel (body) — perfect Victorian gothic feel",
        "Cached glow text rendering (12x faster than naive approach)",
        "Tile-based obsidian texture caching (256x256 master tile)",
        "Fade-to-black screen transitions (0.3s each direction)",
        "Particle effects: dust (explore), eldritch energy (combat), blood splatter",
        "Animated hover: pulsing glow, gold border, shimmer sweep",
        "4-phase victory animation (HP drain → disintegration → pause → fade)",
        "Surface pooling eliminates all per-frame allocations",
        "Custom cursor, side-by-side exploration path cards, typewriter text",
    ]
    for w in working:
        add_bullet(doc, w, color=BONE)

    add_section_header(doc, "Missing / Planned", level=3)
    missing = [
        ("No audio system: ", "The game is completely silent — this is the biggest UX gap"),
        ("No dynamic lighting: ", "HP bars/status icons could emit subtle light"),
        ("Limited vignette: ", "Madness vignette exists but could be more dramatic"),
        ("No character expressions: ", "Sprites are static — no attack/hurt/idle animations"),
        ("No screen transitions variety: ", "All transitions are fade-to-black — could vary by context"),
        ("No background blur on modals: ", "Panels overlay on active background"),
    ]
    for prefix, text in missing:
        add_bullet(doc, text, bold_prefix=prefix, color=OFF_WHITE)

    add_impact_difficulty_table(doc, [
        ("Audio System (SFX + Music)", "★★★★★", "★★★★", "Biggest UX gap — needs AudioManager + assets"),
        ("UI Sound Effects", "★★★★", "★★", "Hover, click, confirm, cancel sounds"),
        ("Dynamic Lighting", "★★★", "★★★", "HP bars/status emit subtle light"),
        ("Character Animations", "★★★", "★★★★", "Attack/hurt/idle states — needs new sprites"),
        ("Context-aware Transitions", "★★", "★★", "Red for danger, purple for eldritch, slide for explore"),
        ("Background Blur on Modals", "★★", "★★", "Progressive blur when panels open"),
        ("Enemy Death Variants", "★★", "★★★", "Multiple death animations per enemy type"),
    ])

    doc.add_page_break()

    # ── 3.5 Content ──
    add_section_header(doc, "3.5  Content & Narrative", level=2)

    add_styled_paragraph(doc, "Status: ⚠️ Needs Expansion — Solid framework, thin content", color=YELLOW, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The content framework is excellent — events, traps, and narratives are data-driven via JSON "
        "files, making expansion trivial. However, the actual content volume is thin: only 6 events, "
        "3 traps, 20 floor narratives, and 10 path templates. Players will see repeats within a single "
        "20-floor run.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Content Inventory", level=3)

    content_items = [
        ("Floor Events", "6", "JSON-driven with outcomes (gold, heal, damage, items, shop access)"),
        ("Traps", "3", "JSON-driven with chance/damage/madness outcomes"),
        ("Floor Narratives", "20", "One per floor — atmospheric text, well-written"),
        ("Path Templates", "10", "Exploration path descriptions with icons"),
        ("Enemy Types", "6 + 1 Boss", "All with unique skills and descriptions"),
        ("Item Templates", "21", "5 weapons, 4 armor, 4 accessories, 4 boots, 4 rings"),
        ("Rarity Tiers", "4", "Common, Uncommon, Rare, Cursed (with stat penalties)"),
    ]

    table = doc.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, ["Category", "Count", "Notes"]):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for cat, count, notes in content_items:
        row_cells = table.add_row().cells
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(cat)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE
        run.font.bold = True

        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(count)
        run.font.size = Pt(9)
        run.font.color.rgb = YELLOW

        set_cell_bg_hex(row_cells[2], "140C26")
        p = row_cells[2].paragraphs[0]
        run = p.add_run(notes)
        run.font.size = Pt(8)
        run.font.color.rgb = OFF_WHITE

    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(10.5)

    add_impact_difficulty_table(doc, [
        ("More Events (6→20+)", "★★★★★", "★★", "JSON-only — just write more event definitions"),
        ("More Traps (3→10+)", "★★★★", "★★", "JSON-only — add variety and trap-specific effects"),
        ("Floor-specific Enemies", "★★★★", "★★★", "New enemy data + unique sprites per floor range"),
        ("Item Set Bonuses", "★★★", "★★★", "Set detection + bonus application in engine"),
        ("More Path Types", "★★★", "★★", "Shrine, altar, portal — add to JSON + screen logic"),
        ("Floor-specific Atmosphere", "★★★", "★★", "Color palette shifts, narrative tone changes per floor"),
        ("Lore Collectibles", "★★", "★★", "Discoverable fragments of Carcosa's history"),
    ])

    doc.add_page_break()

    # ── 3.6 Audio ──
    add_section_header(doc, "3.6  Audio & Polish", level=2)

    add_styled_paragraph(doc, "Status: ❌ Not Implemented — Biggest gap in the game", color=HP_RED, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The game has zero audio. No music, no sound effects, no ambient sound. For a horror-themed "
        "game, this is the single most impactful missing feature. Audio would transform the atmosphere "
        "from 'visual novel' to 'immersive horror experience'. The silence is especially noticeable "
        "during combat, where hits, spells, and enemy attacks have no auditory feedback.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Recommended Audio Layers", level=3)

    audio_layers = [
        ("Music: ", "Ambient drone base + dynamic combat layer that intensifies at low HP"),
        ("UI SFX: ", "Hover (soft chime), click (thud), confirm (bell), cancel (whoosh)"),
        ("Combat SFX: ", "Hit impact, spell cast, block/parry, crit (dramatic), enemy death"),
        ("Environment: ", "Wind, dripping water, distant whispers, creaking wood"),
        ("Madness Audio: ", "Heartbeat layer, whisper overlay, distortion effects at high madness"),
    ]
    for prefix, text in audio_layers:
        add_bullet(doc, text, bold_prefix=prefix, color=OFF_WHITE)

    add_impact_difficulty_table(doc, [
        ("Audio Manager (shared/audio.py)", "★★★★★", "★★★", "Core audio infrastructure — SFX/music/ambient layers"),
        ("UI Sound Effects", "★★★★★", "★", "Low-hanging fruit — immediate feel improvement"),
        ("Combat Sound Effects", "★★★★★", "★★", "Hit, crit, spell, block, death sounds"),
        ("Ambient Soundscapes", "★★★★", "★★", "Per-floor ambient loops"),
        ("Dynamic Music Layers", "★★★", "★★★★", "Music intensity tied to combat/HP state"),
        ("Madness Audio Effects", "★★★", "★★★", "Heartbeat, whispers, distortion at high madness"),
    ])

    # ── 3.7 Save System ──
    add_section_header(doc, "3.7  Save System & Persistence", level=2)

    add_styled_paragraph(doc, "Status: ✅ Working — But limited", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "The save system uses JSON serialization with 5 slots and version 2 format (composed "
        "PlayerIdentity + PlayerProgression + CombatBuffs). It has backward compatibility and "
        "104 tests covering round-trip integrity, version mismatch, corruption handling, and "
        "multi-slot operations.",
        color=OFF_WHITE, space_after=6)

    add_section_header(doc, "Missing Features", level=3)
    save_gaps = [
        ("No meta-progression: ", "No run-to-run unlocks, achievements, or persistent upgrades"),
        ("No auto-save: ", "Manual save only — could auto-save between floors"),
        ("No cloud sync: ", "Local files only"),
        ("No save preview: ", "Can't see class/level/floor before loading"),
    ]
    for prefix, text in save_gaps:
        add_bullet(doc, text, bold_prefix=prefix, color=OFF_WHITE)

    add_impact_difficulty_table(doc, [
        ("Meta-progression System", "★★★★★", "★★★★", "Persistent unlocks across runs — major replayability"),
        ("Auto-save Between Floors", "★★★", "★", "Trivial to implement — just call save on floor advance"),
        ("Save Slot Preview", "★★", "★", "Show class/level/floor in save slot UI"),
        ("Achievement System", "★★★", "★★★", "Track milestones, unlock rewards"),
    ])

    doc.add_page_break()

    # ── 3.8 Testing ──
    add_section_header(doc, "3.8  Testing & CI/CD", level=2)

    add_styled_paragraph(doc, "Status: ✅ Excellent — Industry-leading for a game jam project", color=HP_GREEN, bold=True, space_after=6)

    add_styled_paragraph(doc,
        "With 1,156 automated tests across 5 suites, property-based testing via Hypothesis, "
        "and a GitHub Actions CI pipeline running lint + test across Python 3.10/3.11/3.12, "
        "this project has testing infrastructure that exceeds most commercial game projects.",
        color=OFF_WHITE, space_after=6)

    test_suites = [
        ("test_combat.py", "271", "All classes, damage, buffs, debuffs, AI, boss phases"),
        ("test_edge_cases.py", "161", "0 HP, 0 stats, buff stacking, poison, doom, boss phases"),
        ("test_property_based.py", "515", "Hypothesis-generated random inputs for 10 invariants"),
        ("test_save_load.py", "104", "Round-trip integrity, version mismatch, corruption, multi-slot"),
        ("test_integration.py", "105", "Combat-progression, floor advance, events, DOT, game loop"),
    ]

    table = doc.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, ["Suite", "Tests", "Coverage"]):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for suite, count, coverage in test_suites:
        row_cells = table.add_row().cells
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(suite)
        run.font.size = Pt(9)
        run.font.color.rgb = YELLOW
        run.font.bold = True

        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(count)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE

        set_cell_bg_hex(row_cells[2], "140C26")
        p = row_cells[2].paragraphs[0]
        run = p.add_run(coverage)
        run.font.size = Pt(8)
        run.font.color.rgb = OFF_WHITE

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(10)

    add_impact_difficulty_table(doc, [
        ("Performance Benchmarks", "★★★", "★★", "Benchmark particle systems, glow cache, combat loop"),
        ("UI/Visual Regression Tests", "★★", "★★★", "Screenshot comparison for screen rendering"),
        ("Fuzz Testing for Combat", "★★★", "★★", "Random skill sequences to find edge-case crashes"),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 4. MASTER IMPROVEMENT LIST
    # ═══════════════════════════════════════════

    add_section_header(doc, "4. Master Improvement List", level=1)

    add_styled_paragraph(doc,
        "All improvements sorted by sector, with impact (how much it improves the player experience) "
        "and difficulty (implementation complexity) rated on a 1-5 star scale.",
        color=OFF_WHITE, space_after=10)

    # ── Combat Sector ──
    add_section_header(doc, "⚔️  COMBAT SYSTEM", level=2)

    add_impact_difficulty_table(doc, [
        ("Limb Loss System", "★★★★★", "★★★★", "New engine/limbs.py — limb slots, debuffs, skill restrictions"),
        ("Madness Overhaul (Thresholds)", "★★★★★", "★★★", "Tiers: stable/unstable/delirious/mad — escalating effects"),
        ("Curse System (Limb/Body/Essence)", "★★★★", "★★★", "New engine/curses.py — application, removal, counterplay"),
        ("More Enemy Types (6→12+)", "★★★★", "★★", "Data-driven — add JSON entries + sprite assets"),
        ("Second Boss (Azathoth / Nyarlathotep)", "★★★★", "★★★", "New boss with unique phase mechanics"),
        ("Third Boss (The Pallid Mask)", "★★★", "★★★", "Thematic boss tied to Scholar's lore"),
        ("Enemy Attack Telegraphing (Visual)", "★★★", "★★", "Wind-up animation before powerful attacks"),
        ("Flee Rework (Chase Mechanics)", "★★", "★★", "Partial escape, pursuit consequences"),
        ("Damage Type Resistances", "★★", "★★", "Enemies weak/resistant to specific damage types"),
    ])

    # ── Classes Sector ──
    add_section_header(doc, "🗡️  CLASSES & SKILLS", level=2)

    add_impact_difficulty_table(doc, [
        ("Skill Evolution / Mastery System", "★★★★★", "★★★★", "Skills level up with use — adds long-term progression"),
        ("Skill Branching (Upgrade Paths)", "★★★★", "★★★", "Choose between 2 paths at certain skill levels"),
        ("Class-specific Madness Interactions", "★★★★", "★★★", "Mad Prophet benefits, Scholar knowledge at a cost"),
        ("Starting Loadout Choice", "★★★", "★", "Pick 2 of 4 starting skills for replayability"),
        ("Class Unlock via Achievements", "★★★", "★★", "Not all classes available from start"),
        ("6th Class (e.g., Necromancer)", "★★★", "★★★★", "New class with unique mechanics (summoning?)"),
        ("Passive Skill Trees", "★★★", "★★★★", "Permanent upgrades per class, unlocked through play"),
    ])

    # ── Visual Sector ──
    add_section_header(doc, "🎨  VISUAL & UI", level=2)

    add_impact_difficulty_table(doc, [
        ("Audio System (Full)", "★★★★★", "★★★★", "Music + SFX + ambient — transforms the experience"),
        ("UI Sound Effects", "★★★★★", "★", "Hover, click, confirm — immediate feel improvement"),
        ("Combat Sound Effects", "★★★★★", "★★", "Hit, crit, spell, block, death sounds"),
        ("Dynamic Lighting", "★★★", "★★★", "HP bars/status emit subtle light, torch flicker"),
        ("Character Animations", "★★★", "★★★★", "Attack/hurt/idle states with new sprites"),
        ("Screen Transitions Variety", "★★", "★★", "Context-aware: red danger, purple eldritch, slide explore"),
        ("Background Blur on Modals", "★★", "★★", "Progressive blur when panels open"),
        ("Enemy Presence Effects", "★★", "★★★", "Boss aura distorts screen edges"),
    ])

    doc.add_page_break()

    # ── Content Sector ──
    add_section_header(doc, "📜  CONTENT & NARRATIVE", level=2)

    add_impact_difficulty_table(doc, [
        ("More Events (6→20+)", "★★★★★", "★★", "JSON-only — write more event definitions with varied outcomes"),
        ("More Traps (3→10+)", "★★★★", "★★", "JSON-only — add trap-specific effects and descriptions"),
        ("Floor-specific Enemies", "★★★★", "★★★", "Unique enemies per floor range with themed abilities"),
        ("Item Set Bonuses", "★★★", "★★★", "Equip 2/3 pieces of a set for bonus effects"),
        ("More Path Types (Shrine, Altar)", "★★★", "★★", "New exploration choices with unique mechanics"),
        ("Floor Atmosphere Shifts", "★★★", "★★", "Color palette + narrative tone changes per floor range"),
        ("Lore Collectibles", "★★", "★★", "Discoverable Carcosa lore fragments"),
        ("Random Encounter Chains", "★★", "★★★", "Multi-floor story arcs that remember your choices"),
    ])

    # ── Persistence Sector ──
    add_section_header(doc, "💾  PERSISTENCE & META", level=2)

    add_impact_difficulty_table(doc, [
        ("Meta-progression (Run-to-Run)", "★★★★★", "★★★★", "Persistent unlocks, upgrades across runs"),
        ("Achievement System", "★★★", "★★★", "Track milestones, unlock rewards/classes/items"),
        ("Auto-save Between Floors", "★★★", "★", "Trivial — call save on floor advance"),
        ("Save Slot Preview", "★★", "★", "Show class/level/floor in save slot UI"),
        ("Run Statistics / History", "★★", "★★", "Track runs, deaths, best floors, favorite class"),
        ("Difficulty Modes", "★★★", "★★", "Easy/Normal/Hard with stat multipliers"),
        ("Daily Challenge Seeds", "★★", "★★", "Seeded runs for leaderboard competition"),
    ])

    # ── Technical Sector ──
    add_section_header(doc, "🔧  TECHNICAL & INFRASTRUCTURE", level=2)

    add_impact_difficulty_table(doc, [
        ("Guard Division-by-Zero (4 spots)", "★★★★", "★", "30 min — prevents rare crashes"),
        ("Add Docstrings (26%→80%)", "★★★", "★★★", "3-4 hours — critical for maintainability"),
        ("Rename camelCase Helpers", "★★", "★★", "1 hour — improves consistency"),
        ("Add TypedDict for Data Structures", "★★★", "★★", "2 hours — fixes 17 mypy warnings"),
        ("Performance Benchmarks", "★★★", "★★", "Benchmark particles, glow cache, combat loop"),
        ("Consider ruff as Linter", "★", "★", "Faster alternative to flake8"),
        ("Mod Support / Plugin System", "★★", "★★★★", "Allow community content via JSON mods"),
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 5. PRIORITY ROADMAP
    # ═══════════════════════════════════════════

    add_section_header(doc, "5. Priority Roadmap", level=1)

    add_styled_paragraph(doc,
        "Recommended implementation order based on impact-to-effort ratio. "
        "Each phase builds on the previous one.",
        color=OFF_WHITE, space_after=10)

    # Phase 1
    add_section_header(doc, "Phase 1: Quick Wins (1-2 days)", level=2)
    phase1 = [
        "UI sound effects — hover, click, confirm (★★★★★ impact, ★ difficulty)",
        "Auto-save between floors (★★★ impact, ★ difficulty)",
        "Guard division-by-zero in 4 spots (★★★★ impact, ★ difficulty)",
        "Save slot preview showing class/level/floor (★★ impact, ★ difficulty)",
        "More events and traps via JSON (★★★★★ impact, ★★ difficulty)",
    ]
    for item in phase1:
        add_bullet(doc, item, color=BONE)

    # Phase 2
    add_section_header(doc, "Phase 2: Combat Depth (1-2 weeks)", level=2)
    phase2 = [
        "Madness overhaul — threshold system with escalating effects (★★★★★ impact, ★★★ difficulty)",
        "More enemy types — expand from 6 to 12+ (★★★★ impact, ★★ difficulty)",
        "Second boss encounter (★★★★ impact, ★★★ difficulty)",
        "Combat sound effects (★★★★★ impact, ★★ difficulty)",
        "Limb loss system — new engine/limbs.py (★★★★★ impact, ★★★★ difficulty)",
    ]
    for item in phase2:
        add_bullet(doc, item, color=BONE)

    # Phase 3
    add_section_header(doc, "Phase 3: Progression & Polish (2-4 weeks)", level=2)
    phase3 = [
        "Audio system — full music + ambient + SFX infrastructure (★★★★★ impact, ★★★★ difficulty)",
        "Skill evolution / mastery — skills level with use (★★★★★ impact, ★★★★ difficulty)",
        "Curse system — limb/body/essence curses (★★★★ impact, ★★★ difficulty)",
        "Meta-progression — run-to-run unlocks (★★★★★ impact, ★★★★ difficulty)",
        "Item set bonuses (★★★ impact, ★★★ difficulty)",
    ]
    for item in phase3:
        add_bullet(doc, item, color=BONE)

    # Phase 4
    add_section_header(doc, "Phase 4: Endgame & Community (1-2 months)", level=2)
    phase4 = [
        "Eight-God Religion system — faith tracking, god-specific mechanics (★★★★ impact, ★★★★ difficulty)",
        "Skill branching — upgrade paths at certain levels (★★★★ impact, ★★★ difficulty)",
        "Achievement system (★★★ impact, ★★★ difficulty)",
        "Dynamic lighting and character animations (★★★ impact, ★★★★ difficulty)",
        "Mod support — allow community content via JSON plugins (★★ impact, ★★★★ difficulty)",
    ]
    for item in phase4:
        add_bullet(doc, item, color=BONE)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 6. APPENDIX
    # ═══════════════════════════════════════════

    add_section_header(doc, "6. Appendix: Technical Metrics", level=1)

    add_section_header(doc, "Code Quality Scorecard", level=2)

    scorecard = [
        ("Architecture", "A", "Clean layered design, DI, well-organized packages"),
        ("Test Coverage", "A+", "1,156 tests, property-based testing, CI/CD pipeline"),
        ("Combat System", "A", "Deep mechanics, 67+ buffs, data-driven registries"),
        ("Visual Design", "A", "Cohesive Lovecraftian aesthetic, particle effects, animations"),
        ("Content Volume", "C+", "Framework is great, but 6 events and 3 traps is thin"),
        ("Audio", "F", "Not implemented — silence throughout"),
        ("Save System", "B+", "Working with tests, but no meta-progression"),
        ("Documentation", "B+", "ROADMAP, GAME_MEMORY, CLASS_REFERENCE, DEVELOPER_GUIDE all present"),
        ("Code Style", "B", "Good overall, 31 camelCase helpers, 26% docstring coverage"),
        ("Performance", "A−", "Surface pooling, texture caching, glow text cache"),
    ]

    table = doc.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, ["Area", "Grade", "Notes"]):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for area, grade, notes in scorecard:
        row_cells = table.add_row().cells
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(area)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE
        run.font.bold = True

        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(grade)
        run.font.size = Pt(10)
        run.font.color.rgb = HP_GREEN if grade.startswith("A") else YELLOW if grade.startswith("B") else HP_RED
        run.font.bold = True

        set_cell_bg_hex(row_cells[2], "140C26")
        p = row_cells[2].paragraphs[0]
        run = p.add_run(notes)
        run.font.size = Pt(8)
        run.font.color.rgb = OFF_WHITE

    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(11)

    add_styled_paragraph(doc, "", font_size=12, space_after=12)

    add_section_header(doc, "Color Palette Reference", level=2)

    colors = [
        ("VOID", "1A0A2E", "Deep dark purple — page backgrounds"),
        ("SHADOW", "2D1B4E", "Dark purple — panel backgrounds"),
        ("ELDRITCH", "4A0E4E", "Purple — eldritch energy effects"),
        ("YELLOW", "D4A017", "Gold — primary accent, UI highlights"),
        ("AMBER", "B8860B", "Dark gold — secondary accent"),
        ("BONE", "D4C5A9", "Light beige — body text"),
        ("CRIMSON", "8B0000", "Dark red — damage, danger"),
        ("GOLD", "F1C40F", "Bright gold — titles, rewards"),
        ("ELDRITCH_GOLD", "E8B92D", "Gold — headers, dividers"),
        ("ELDRITCH_PURPLE", "AF82E1", "Light purple — secondary text"),
        ("HP_GREEN", "2ECC71", "Green — health, positive effects"),
        ("HP_RED", "E74C3C", "Red — low health, warnings"),
    ]

    table = doc.add_table(rows=1, cols=4)
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, ["Name", "Hex", "Sample", "Usage"]):
        set_cell_bg_hex(cell, "1A0A2E")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.color.rgb = ELDRITCH_GOLD
        run.font.bold = True

    for name, hex_val, usage in colors:
        row_cells = table.add_row().cells
        set_cell_bg_hex(row_cells[0], "140C26")
        p = row_cells[0].paragraphs[0]
        run = p.add_run(name)
        run.font.size = Pt(9)
        run.font.color.rgb = BONE
        run.font.bold = True

        set_cell_bg_hex(row_cells[1], "140C26")
        p = row_cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"#{hex_val}")
        run.font.size = Pt(9)
        run.font.color.rgb = OFF_WHITE

        set_cell_bg_hex(row_cells[2], hex_val)
        p = row_cells[2].paragraphs[0]
        run = p.add_run("██████")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(hex_val)

        set_cell_bg_hex(row_cells[3], "140C26")
        p = row_cells[3].paragraphs[0]
        run = p.add_run(usage)
        run.font.size = Pt(8)
        run.font.color.rgb = OFF_WHITE

    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(8.5)

    add_styled_paragraph(doc, "", font_size=12, space_after=20)

    # Final note
    add_divider(doc, GOLD_TRIM)
    add_styled_paragraph(doc,
        "⛧ The Yellow Sign fades. For now. ⛧",
        font_size=12, color=ELDRITCH_PURPLE, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

    # ── Save ──
    output_path = os.path.join(os.path.dirname(__file__), "Yellow_Sign_Analysis.docx")
    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_analysis_document()
